from typing import Tuple
from pypdf import PdfReader
from docx import Document as DocxDocument

def read_text_from_pdf(file_path: str) -> str:
    pdf = PdfReader(file_path)
    parts = []
    for page in pdf.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)

def read_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file including both paragraphs and tables.
    Tables are converted to readable text format.
    """
    doc = DocxDocument(file_path)
    parts = []
    
    # Extract all paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    
    # Extract all tables
    for table in doc.tables:
        table_text = extract_table_text(table)
        if table_text:
            parts.append("\n" + table_text)
    
    return "\n\n".join(parts)

def extract_table_text(table) -> str:
    """
    Convert a DOCX table to readable text format.
    Each row is preserved with clear separators.
    """
    lines = []
    
    for i, row in enumerate(table.rows):
        # Get cell values
        cells = [cell.text.strip() for cell in row.cells]
        
        # Skip completely empty rows
        if not any(cells):
            continue
        
        # Join cells with pipe separator for readability
        row_text = " | ".join(cells)
        lines.append(row_text)
    
    return "\n".join(lines)

def read_text_from_txt(file_path: str, encoding="utf-8") -> str:
    with open(file_path, "r", encoding=encoding, errors="ignore") as f:
        return f.read()

def read_any(file_path: str, mime: str, filename: str) -> Tuple[str, str]:
    name = filename.lower()
    if name.endswith(".pdf") or mime == "application/pdf":
        return read_text_from_pdf(file_path), "pdf"
    if name.endswith(".docx") or mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
        return read_text_from_docx(file_path), "docx"
    # default to txt
    return read_text_from_txt(file_path), "txt"

def simple_chunks(text: str, target_chars: int = 1200, overlap: int = 150):
    """
    Split text into chunks with overlap.
    Improved to avoid creating too many tiny chunks.
    Preserves line breaks for better structure.
    """
    # Normalize spaces within lines but preserve line breaks
    lines = text.split('\n')
    normalized_lines = [' '.join(line.split()) for line in lines]
    text = '\n'.join(normalized_lines)
    
    n = len(text)
    
    # If text is shorter than target, return as single chunk
    if n <= target_chars:
        yield text
        return
    
    i = 0
    while i < n:
        # Calculate end position
        j = min(i + target_chars, n)
        
        # If we're near the end, just take the rest
        if j >= n - overlap:
            yield text[i:].strip()
            break
        
        # Try to cut on paragraph boundary (double newline)
        k = text.rfind('\n\n', i, j)
        if k != -1:
            k += 2
        else:
            # Try to cut on single newline
            k = text.rfind('\n', i, j)
            if k != -1:
                k += 1
            else:
                # Try to cut on sentence end
                k = text.rfind('. ', i, j)
                if k != -1 and j - k <= 180:
                    k += 2
                else:
                    # Cut at word boundary
                    k = text.rfind(' ', i, j)
                    if k == -1 or j - k > 180:
                        k = j
        
        chunk = text[i:k].strip()
        if chunk:  # Only yield non-empty chunks
            yield chunk
        
        # Move forward, accounting for overlap
        i = max(k - overlap, i + 1)
        
        # Safety check: ensure we're making progress
        if i >= n:
            break
